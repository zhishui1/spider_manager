"""
flkgov 数据库连接和初始化脚本
"""
import re
import psycopg2
import os
from datetime import datetime
import json
from pathlib import Path
from docx import Document
from pathlib import Path
import json
import re
import os


def read_doc_file(file_path: str) -> str:
    """读取doc或docx文件内容"""
    file_path_obj = Path(file_path)
    try:
        if file_path_obj.suffix.lower() == '.docx':
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif file_path_obj.suffix.lower() == '.doc':
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = False
                try:
                    doc = word.Documents.Open(str(file_path_obj.absolute()))
                    text = doc.Content.Text
                    doc.Close(False)
                    return text
                finally:
                    word.Quit()
            except Exception:
                return ""
        return ""
    except Exception:
        return ""


DB_CONFIG = {
    'host': '192.168.1.40',
    'port': 5432,
    'database': 'flkgov_db',
    'user': 'flkgov_user',
    'password': '1421nbnb'
}

DATA_DIR = Path(r"E:\spider_manager\data\flkgov")
JSON_FILE = DATA_DIR / "flkgov_data.json"
FILES_DIR = DATA_DIR / "flkgov_files"


import re


def parse_docx_text(full_text: str) -> dict:
    """解析文档内容，提取摘要、章节和条目"""
    abstract = ""
    chapters = []
    current_chapter_no = None
    current_articles = []
    
    lines = full_text.split("\n")
    
    chapter_pattern = re.compile(r'^(第[一二三四五六七八九十百千零\d]+章)(.*)')
    article_pattern = re.compile(r'^(第[一二三四五六七八九十百千零\d]+条)(.*)')
    enum_pattern = re.compile(r'^（[一二三四五六七八九十百千零\d]+）')
    
    def process_enum_newlines(text: str) -> str:
        """处理枚举项，在每个（n）前添加一个换行"""
        result = []
        i = 0
        while i < len(text):
            if text[i] == '（':
                if result and result[-1] != '\n':
                    result.append('\n')
                result.append('（')
                i += 1
            else:
                result.append(text[i])
                i += 1
        return ''.join(result)
    
    in_abstract = True
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        match_chapter = chapter_pattern.match(line)
        match_article = article_pattern.match(line)
        
        if match_chapter:
            if current_chapter_no and current_articles:
                chapters.append({
                    'chapter_no': current_chapter_no,
                    'articles': current_articles
                })
            in_abstract = False
            chapter_title = match_chapter.group(1)
            chapter_name = match_chapter.group(2).strip() if match_chapter.group(2) else ""
            if chapter_name:
                current_chapter_no = f"{chapter_title} - {chapter_name}"
            else:
                current_chapter_no = chapter_title
            current_articles = []
        elif match_article and in_abstract:
            in_abstract = False
            current_chapter_no = '第一章'
            current_articles = []
            content = match_article.group(2).strip() if match_article.group(2) else ""
            current_articles.append({
                'article_no': match_article.group(1),
                'content': content
            })
        elif in_abstract:
            abstract += line + "\n"
        else:
            if match_article:
                if current_articles:
                    current_articles[-1]['content'] = process_enum_newlines(current_articles[-1]['content'].strip())
                content = match_article.group(2).strip() if match_article.group(2) else ""
                current_articles.append({
                    'article_no': match_article.group(1),
                    'content': content
                })
            else:
                if current_articles:
                    current_articles[-1]['content'] += line + "\n"
    
    if current_articles:
        current_articles[-1]['content'] = process_enum_newlines(current_articles[-1]['content'].strip())
        if current_chapter_no:
            chapters.append({
                'chapter_no': current_chapter_no,
                'articles': current_articles
            })
    
    if not chapters and abstract:
        chapters.append({
            'chapter_no': '第一章',
            'articles': []
        })
    
    abstract = abstract.strip()
    
    return {
        'abstract': abstract,
        'chapters': chapters
    }


def read_json_and_docx():
    """按行读取flkgov_data.json，提取字段并读取对应的docx文件内容"""
    if not JSON_FILE.exists():
        print(f"JSON文件不存在: {JSON_FILE}")
        return
    
    if not FILES_DIR.exists():
        print(f"文件目录不存在: {FILES_DIR}")
        return
    
    print(f"正在读取: {JSON_FILE}")
    print(f"文件目录: {FILES_DIR}")
    print("=" * 60)
    
    count = 0
    with open(JSON_FILE, 'r', encoding='utf-8') as f:
        for line_num, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            
            try:
                item = json.loads(line)
                
                item_id = item.get('item_id')
                title = item.get('title')
                url = item.get('url')
                publish_date = item.get('publish_date')
                
                data = item.get('data', {})
                category = data.get('category')
                regulation_type = data.get('regulation_type')
                issuing_body = data.get('issuing_body')
                promulgation_date = data.get('promulgation_date')
                effective_date = data.get('effective_date')
                status_code = data.get('status_code')
                status = data.get('status')
                
                print(f"\n{'='*60}")
                print(f"第 {line_num} 条记录")
                print(f"{'='*60}")
                print(f"item_id: {item_id}")
                print(f"title: {title}")
                print(f"url: {url}")
                print(f"category: {category}")
                print(f"regulation_type: {regulation_type}")
                print(f"issuing_body: {issuing_body}")
                print(f"promulgation_date: {promulgation_date}")
                print(f"effective_date: {effective_date}")
                print(f"status: {status}")
                
                item_files_dir = FILES_DIR / str(item_id)
                if item_files_dir.exists():
                    docx_files = list(item_files_dir.glob("*.docx")) + list(item_files_dir.glob("*.doc"))
                    
                    if docx_files:
                        docx_file = docx_files[0]
                        print(f"\n--- 解析文件: {docx_file.name} ---")
                        
                        try:
                            full_text = read_doc_file(str(docx_file))
                            parsed = parse_docx_text(full_text)
                            
                            print(f"\n【摘要】(从开始到第一章之间)")
                            print("-" * 40)
                            print(parsed['abstract'] if parsed['abstract'] else "(无)")
                            
                            print(f"\n【章节和条款】")
                            print("-" * 40)
                            
                            for chapter in parsed['chapters']:
                                print(f"\n章节: {chapter['chapter_no']}")
                                for i, article in enumerate(chapter['articles']):
                                    is_last = (i == len(chapter['articles']) - 1)
                                    print(f"\n  条款: {article['article_no']}" + (" (本章最后一条)" if is_last else ""))
                                    content = article['content'].strip()
                                    if content:
                                        print(f"  内容:\n{content}")
                            
                        except Exception as e:
                            print(f"解析文件失败: {e}")
                else:
                    print(f"文件目录不存在: {item_files_dir}")
                
                count += 1
                if count >= 1:
                    print("\n" + "=" * 60)
                    print(f"测试完成")
                    break
                    
            except json.JSONDecodeError as e:
                print(f"JSON解析错误 (行 {line_num}): {e}")
                continue


def get_connection():
    """获取数据库连接"""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        return conn
    except Exception as e:
        print(f"连接数据库失败: {e}")
        return None


def init_tables():
    """初始化数据库表（删除旧表后重新创建）"""
    conn = get_connection()
    if not conn:
        return False
    
    cursor = conn.cursor()
    
    # 删除旧表（先删子表，再删主表）
    cursor.execute("DROP TABLE IF EXISTS flkgov_articles CASCADE;")
    cursor.execute("DROP TABLE IF EXISTS flkgov_regulations CASCADE;")
    conn.commit()
    print("已删除旧表")
    
    # 创建法规信息表
    cursor.execute("""
        CREATE TABLE flkgov_regulations (
            id SERIAL PRIMARY KEY,
            item_id VARCHAR(50),
            url VARCHAR(500),
            title VARCHAR(500),
            regulation_type VARCHAR(100),
            issuing_body VARCHAR(200),
            promulgation_date DATE,
            effective_date DATE,
            status VARCHAR(20),
            abstract TEXT,
            chapter_num INT DEFAULT 0,
            article_num INT DEFAULT 0
        );
    """)
    
    # 创建索引
    cursor.execute("CREATE INDEX idx_regulations_title ON flkgov_regulations(title);")
    cursor.execute("CREATE INDEX idx_regulations_issuing_body ON flkgov_regulations(issuing_body);")
    
    # 创建条目信息表
    cursor.execute("""
        CREATE TABLE flkgov_articles (
            id SERIAL PRIMARY KEY,
            regulation_id INT REFERENCES flkgov_regulations(id),
            chapter_no VARCHAR(200),
            article_no VARCHAR(50),
            content TEXT
        );
    """)
    
    # 创建索引
    cursor.execute("CREATE INDEX idx_articles_regulation_id ON flkgov_articles(regulation_id);")
    
    conn.commit()
    cursor.close()
    conn.close()
    
    print("数据库表初始化完成!")
    return True


def test_connection():
    """测试数据库连接"""
    print("=" * 50)
    print("测试 PostgreSQL 数据库连接")
    print("=" * 50)
    print(f"Host: {DB_CONFIG['host']}")
    print(f"Port: {DB_CONFIG['port']}")
    print(f"Database: {DB_CONFIG['database']}")
    print(f"User: {DB_CONFIG['user']}")
    print("=" * 50)
    
    conn = get_connection()
    if conn:
        print("✓ 数据库连接成功!")
        
        cursor = conn.cursor()
        cursor.execute("SELECT version();")
        version = cursor.fetchone()
        print(f"\nPostgreSQL 版本: {version[0]}")
        
        # 查看所有表
        cursor.execute("""
            SELECT table_name 
            FROM information_schema.tables 
            WHERE table_schema = 'public'
            ORDER BY table_name;
        """)
        tables = cursor.fetchall()
        
        print("\n当前数据库中的表:")
        for table in tables:
            print(f"  - {table[0]}")
        
        cursor.close()
        conn.close()
        return True
    else:
        print("✗ 数据库连接失败!")
        return False


def insert_regulation(data: dict) -> int:
    """插入法规记录，返回ID"""
    conn = get_connection()
    if not conn:
        return None
    
    cursor = conn.cursor()
    
    cursor.execute("""
        INSERT INTO flkgov_regulations (
            item_id, url, title, regulation_type, issuing_body,
            promulgation_date, effective_date, status, abstract,
            chapter_num, article_num
        ) VALUES (
            %(item_id)s, %(url)s, %(title)s, %(regulation_type)s, %(issuing_body)s,
            %(promulgation_date)s, %(effective_date)s, %(status)s, %(abstract)s,
            %(chapter_num)s, %(article_num)s
        )
        ON CONFLICT (item_id) DO UPDATE SET
            title = EXCLUDED.title,
            url = EXCLUDED.url,
            regulation_type = EXCLUDED.regulation_type,
            issuing_body = EXCLUDED.issuing_body,
            promulgation_date = EXCLUDED.promulgation_date,
            effective_date = EXCLUDED.effective_date,
            status = EXCLUDED.status,
            abstract = EXCLUDED.abstract,
            chapter_num = EXCLUDED.chapter_num,
            article_num = EXCLUDED.article_num
        RETURNING id;
    """, {
        'item_id': data.get('item_id'),
        'url': data.get('url'),
        'title': data.get('title'),
        'regulation_type': data.get('regulation_type'),
        'issuing_body': data.get('issuing_body'),
        'promulgation_date': data.get('promulgation_date'),
        'effective_date': data.get('effective_date'),
        'status': data.get('status'),
        'abstract': data.get('abstract'),
        'chapter_num': data.get('chapter_num', 0),
        'article_num': data.get('article_num', 0)
    })
    
    regulation_id = cursor.fetchone()[0]
    conn.commit()
    cursor.close()
    conn.close()
    
    return regulation_id


def insert_article(regulation_id: int, data: dict):
    """插入条目记录"""
    conn = get_connection()
    if not conn:
        return False
    
    cursor = conn.cursor()
    
    cursor.execute("""
        INSERT INTO flkgov_articles (
            regulation_id, chapter_no, article_no, content
        ) VALUES (
            %(regulation_id)s, %(chapter_no)s, %(article_no)s, %(content)s
        )
    """, {
        'regulation_id': regulation_id,
        'chapter_no': data.get('chapter_no'),
        'article_no': data.get('article_no'),
        'content': data.get('content')
    })
    
    conn.commit()
    cursor.close()
    conn.close()
    
    return True


def read_first_and_last_3():
    """读取前3条和后3条法规，验证解析结果"""
    if not JSON_FILE.exists():
        print(f"JSON文件不存在: {JSON_FILE}")
        return
    
    if not FILES_DIR.exists():
        print(f"文件目录不存在: {FILES_DIR}")
        return
    
    with open(JSON_FILE, 'r', encoding='utf-8') as f:
        all_lines = f.readlines()
    
    print(f'总记录数: {len(all_lines)}')
    
    def process_item(line):
        item = json.loads(line)
        item_id = item.get('item_id')
        title = item.get('title')
        url = item.get('url')
        
        data = item.get('data', {})
        category = data.get('category')
        regulation_type = data.get('regulation_type')
        issuing_body = data.get('issuing_body')
        promulgation_date = data.get('promulgation_date')
        effective_date = data.get('effective_date')
        status = data.get('status')
        
        result = {
            'item_id': item_id,
            'title': title,
            'url': url,
            'category': category,
            'regulation_type': regulation_type,
            'issuing_body': issuing_body,
            'promulgation_date': promulgation_date,
            'effective_date': effective_date,
            'status': status
        }
        
        item_files_dir = FILES_DIR / str(item_id)
        if item_files_dir.exists():
            docx_files = list(item_files_dir.glob("*.docx")) + list(item_files_dir.glob("*.doc"))
            if docx_files:
                full_text = read_doc_file(str(docx_files[0]))
                parsed = parse_docx_text(full_text)
                result['abstract'] = parsed['abstract']
                result['chapters'] = parsed['chapters']
                result['article_count'] = sum(len(ch['articles']) for ch in parsed['chapters'])
        
        return result
    
    print('\n' + '='*60)
    print('前3条法规')
    print('='*60)
    
    for i in range(3):
        item = process_item(all_lines[i])
        print(f'\n--- 第{i+1}条 ---')
        print(f"item_id: {item['item_id']}")
        print(f"title: {item['title']}")
        print(f"category: {item['category']}")
        print(f"regulation_type: {item['regulation_type']}")
        print(f"issuing_body: {item['issuing_body']}")
        print(f"promulgation_date: {item['promulgation_date']}")
        print(f"effective_date: {item['effective_date']}")
        print(f"status: {item['status']}")
        if 'abstract' in item:
            print(f"abstract长度: {len(item['abstract'])} 字符")
            print(f"章节数: {len(item['chapters'])}")
            print(f"条款数: {item['article_count']}")
    
    print('\n' + '='*60)
    print('后3条法规')
    print('='*60)
    
    for i in range(3):
        item = process_item(all_lines[-(i+1)])
        print(f'\n--- 第{len(all_lines)-i}条 ---')
        print(f"item_id: {item['item_id']}")
        print(f"title: {item['title']}")
        print(f"category: {item['category']}")
        print(f"regulation_type: {item['regulation_type']}")
        print(f"issuing_body: {item['issuing_body']}")
        print(f"promulgation_date: {item['promulgation_date']}")
        print(f"effective_date: {item['effective_date']}")
        print(f"status: {item['status']}")
        if 'abstract' in item:
            print(f"abstract长度: {len(item['abstract'])} 字符")
            print(f"章节数: {len(item['chapters'])}")
            print(f"条款数: {item['article_count']}")
    
    print('\n测试完成')


if __name__ == '__main__':
    import sys
    
    if len(sys.argv) > 1:
        if sys.argv[1] == 'read':
            read_json_and_docx()
        elif sys.argv[1] == 'read6':
            read_first_and_last_3()
        elif sys.argv[1] == 'test':
            test_connection()
        elif sys.argv[1] == 'init':
            init_tables()
        else:
            print("未知命令")
    else:
        if test_connection():
            init_tables()
